from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("/Users/Admin1/Documents/Codex/2026-07-09/xian/字节财经业务_AI全栈实习_面试经验帖.docx")

FONT_ASCII = "Calibri"
# Use a CJK font that is present in the document-rendering environment.
FONT_CJK = "Hiragino Sans GB"
BLACK = RGBColor(0x22, 0x22, 0x22)
NAVY = RGBColor(0x20, 0x37, 0x48)
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = "F4F6F9"


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = FONT_ASCII
    run._element.get_or_add_rPr()
    fonts = run._element.rPr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), FONT_ASCII)
    fonts.set(qn("w:hAnsi"), FONT_ASCII)
    fonts.set(qn("w:eastAsia"), FONT_CJK)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size, color=BLACK, bold=False):
    style.font.name = FONT_ASCII
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    fonts = rpr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), FONT_ASCII)
    fonts.set(qn("w:hAnsi"), FONT_ASCII)
    fonts.set(qn("w:eastAsia"), FONT_CJK)


def set_cell_shading_or_paragraph_shading(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_left_border(paragraph, color="2E74B5", size="18", space="8"):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), space)
    left.set(qn("w:color"), color)
    pbdr.append(left)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=GRAY)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, display, fld_end])
    run2 = paragraph.add_run(" 页")
    set_run_font(run2, size=9, color=GRAY)


def add_numbering_definition(doc, num_fmt, lvl_text, left_twips, hanging_twips, font=None):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(x.get(qn("w:abstractNumId")))
        for x in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), num_fmt)
    lvl.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), lvl_text)
    lvl.append(text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left_twips))
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left_twips))
    ind.set(qn("w:hanging"), str(hanging_twips))
    ppr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "290")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    lvl.append(ppr)
    if font:
        rpr = OxmlElement("w:rPr")
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), font)
        rfonts.set(qn("w:hAnsi"), font)
        rpr.append(rfonts)
        lvl.append(rpr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    numpr.append(ilvl)
    numpr.append(numid)


def add_bullet(doc, text, bullet_num_id):
    p = doc.add_paragraph(style="List Body")
    apply_numbering(p, bullet_num_id)
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    run = p.add_run(text)
    set_run_font(run, size=11, color=BLACK)
    return p


def add_numbered(doc, text, decimal_num_id):
    p = doc.add_paragraph(style="List Body")
    apply_numbering(p, decimal_num_id)
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    run = p.add_run(text)
    set_run_font(run, size=11, color=BLACK)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="Normal")
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        set_run_font(run, bold=True)
        run2 = p.add_run(text[len(bold_prefix):])
        set_run_font(run2)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_question(doc, question, answer, reflection=None):
    p = doc.add_paragraph(style="Question")
    r = p.add_run(question)
    set_run_font(r, size=11.5, color=DARK_BLUE, bold=True)
    add_body(doc, f"我的回答：{answer}", bold_prefix="我的回答：")
    if reflection:
        add_body(doc, f"复盘：{reflection}", bold_prefix="复盘：")


def add_callout(doc, text, label="一句话总结"):
    p = doc.add_paragraph(style="Callout")
    set_cell_shading_or_paragraph_shading(p, LIGHT_GRAY)
    set_paragraph_left_border(p)
    p.paragraph_format.left_indent = Inches(0.14)
    p.paragraph_format.right_indent = Inches(0.14)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.25
    r1 = p.add_run(f"{label}｜")
    set_run_font(r1, size=11, color=DARK_BLUE, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=11, color=BLACK)
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)
section.different_first_page_header_footer = True

styles = doc.styles
normal = styles["Normal"]
set_style_font(normal, 11, BLACK)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.333

for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
]:
    st = styles[name]
    set_style_font(st, size, color, bold=True)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True
    st.paragraph_format.keep_together = True

if "List Body" in [s.name for s in styles]:
    list_body = styles["List Body"]
else:
    list_body = styles.add_style("List Body", WD_STYLE_TYPE.PARAGRAPH)
    list_body.base_style = styles["Normal"]
set_style_font(list_body, 11, BLACK)
list_body.paragraph_format.space_before = Pt(0)
list_body.paragraph_format.space_after = Pt(4)
list_body.paragraph_format.line_spacing = 1.208

question_style = styles.add_style("Question", WD_STYLE_TYPE.PARAGRAPH)
set_style_font(question_style, 11.5, DARK_BLUE, bold=True)
question_style.paragraph_format.space_before = Pt(8)
question_style.paragraph_format.space_after = Pt(3)
question_style.paragraph_format.keep_with_next = True

callout_style = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
set_style_font(callout_style, 11, BLACK)
callout_style.paragraph_format.space_before = Pt(6)
callout_style.paragraph_format.space_after = Pt(12)
callout_style.paragraph_format.line_spacing = 1.25

bullet_num_id = add_numbering_definition(
    doc, "bullet", "•", left_twips=540, hanging_twips=280, font="Symbol"
)
decimal_num_id = add_numbering_definition(
    doc, "decimal", "%1.", left_twips=540, hanging_twips=280
)
decimal_prepare_num_id = add_numbering_definition(
    doc, "decimal", "%1.", left_twips=540, hanging_twips=280
)

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = hp.add_run("面试经验帖 · 2026.07")
set_run_font(hr, size=9, color=GRAY)

footer = section.footer
fp = footer.paragraphs[0]
add_page_number(fp)

# Editorial cover: a restrained, cover-like opening for a long-form experience post.
cover_spacer = doc.add_paragraph()
cover_spacer.paragraph_format.space_before = Pt(72)
cover_spacer.paragraph_format.space_after = Pt(0)

kicker = doc.add_paragraph()
kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
kicker.paragraph_format.space_after = Pt(16)
kr = kicker.add_run("面 试 复 盘")
set_run_font(kr, size=11, color=BLUE, bold=True)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(8)
tr = title.add_run("非科班大三，第一次被项目细节问穿")
set_run_font(tr, size=28, color=NAVY, bold=True)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(28)
sr = subtitle.add_run("字节跳动财经业务 AI 全栈实习面经")
set_run_font(sr, size=15, color=DARK_BLUE)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_after = Pt(42)
mr = meta.add_run("线上视频面试  ·  约 30 分钟  ·  2026 年 7 月")
set_run_font(mr, size=10.5, color=GRAY)

quote = doc.add_paragraph()
quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
quote.paragraph_format.left_indent = Inches(0.55)
quote.paragraph_format.right_indent = Inches(0.55)
quote.paragraph_format.space_after = Pt(18)
qr = quote.add_run(
    "“能借助 AI 把产品做出来”只是起点；面试真正追问的是："
    "你是否理解系统、能否解释取舍、出了问题能不能自己接住。"
)
set_run_font(qr, size=13, color=GRAY, italic=True)

note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
nr = note.add_run(
    "本文根据会议实时转写整理，已删除口头语并对表达做结构化改写；"
    "个别识别不清的技术名词使用概括表述。"
)
set_run_font(nr, size=9.5, color=GRAY)

doc.add_page_break()

doc.add_heading("先说结论", level=1)
add_body(
    doc,
    "这场面试对我最大的提醒，是“做过项目”和“掌握项目”之间差得很远。"
    "我确实独立拆了需求、推进了一个接近上线的小程序，也用 AI 完成了大量开发工作；"
    "但当面试官继续追问部署环境、轮询机制、缓存、并发抢单和代码审查时，"
    "我有不少细节只能说“做过，但讲不清”。"
)
add_body(
    doc,
    "面试没有刻意为难，问题几乎都围绕我主动介绍的项目展开。"
    "真正暴露短板的不是不会某个框架，而是我对 AI 生成代码的理解和控制还不够，"
    "技术表达也缺少清晰的结构。对非科班、AI 辅助开发型选手来说，这次经历很有代表性。"
)
add_callout(
    doc,
    "项目能运行，不等于我能解释它为什么这样运行；能发现 bug，也不等于我能定位到机制层面的原因。"
)

doc.add_heading("一、我的背景与投递情况", level=1)
add_body(
    doc,
    "我是福州大学化学学院化学专业的大三学生，非计算机科班。"
    "平时接触过 Python、C++、TypeScript 和 JavaScript，前后端都做过一些，"
    "但系统性的计算机基础、前端框架和工程化训练都比较薄弱。"
)
add_body(
    doc,
    "这次是在 BOSS 直聘上看到岗位并投递的。岗位方向与财经业务相关，"
    "描述里提到了 AI 全栈应用开发。收到面试邀请时我其实有点意外，"
    "也因此没有按照传统前端或后端岗位的标准做足准备。"
)

doc.add_heading("二、面试流程与问题复盘", level=1)
add_body(
    doc,
    "整场面试大约 30 分钟，节奏很快。流程基本是：自我介绍 → 项目深挖 →"
    " AI 编程与基础能力追问 → 个人优劣势 → 反向提问。"
)

doc.add_heading("1. 自我介绍：开场就进入项目", level=2)
add_question(
    doc,
    "你会哪些编程语言？做过什么项目？",
    "我介绍了自己会 Python、C++、TypeScript 和 JavaScript，随后重点讲了一个类似同城配送平台的小程序项目。",
    "我的开场信息比较散，技术栈和项目亮点没有形成主线。更好的说法应该是：先交代非科班背景，再用一句话说明项目、我的职责和当前进度，最后点出两个最值得深挖的技术点。",
)

doc.add_heading("2. 项目深挖：从“做了什么”追到“为什么这样做”", level=2)
add_body(
    doc,
    "这是整场面试的核心。我做的是一个同城配送类项目，包含商家 Web 端，"
    "以及合并在微信小程序里的用户端和骑手端。用户可注册并申请成为骑手，"
    "商家审核申请、管理骑手、接收订单并派单，骑手上线后可以查看附近订单并抢单。"
    "项目来自真实委托，目前主要功能已完成，正式上线还在等待营业主体、域名和平台审核等条件。"
)

add_question(
    doc,
    "这个项目是谁的需求？你具体负责什么？",
    "项目来自熟人介绍的真实商业需求。我负责拆需求、梳理业务流程、设计多端形态，并借助 AI 完成开发、联调和手动测试。",
    "“真实需求”是加分项，但我没有马上讲清楚自己的责任边界。如果项目大量由 AI 写代码，更应该准确区分：哪些决策是我做的、哪些实现由 AI 完成、哪些关键模块由我亲自验证。",
)

add_question(
    doc,
    "整体技术方案是什么？",
    "后端基于 NestJS，业务数据存储在 PostgreSQL 中，地理位置能力使用 PostGIS，并接入腾讯地图服务；此外还使用缓存处理部分短生命周期信息和订单状态。项目部署在第三方云平台上。",
    "原始回答里框架、数据库、地图和部署信息是逐个蹦出来的，缺少架构图式的表达。面试时应按“客户端—服务端—数据库—缓存—第三方服务—部署”六层来讲。",
)

add_question(
    doc,
    "后端和数据库怎么部署？运行在什么环境？",
    "我能确认后端和数据库已经部署成功，也配置过域名、API 和备份，但具体环境和部署步骤主要是跟着 AI 一步步完成的，现场没能完整复述。",
    "这是第一个明显失分点。只要简历上写了“部署”，至少要能说清运行时、构建与启动方式、环境变量、数据库连接、域名与 HTTPS、日志、备份和故障排查。",
)

add_question(
    doc,
    "缓存用来做什么？",
    "我回答它用于保存短生命周期的信息，以及与订单状态有关的数据。",
    "这个回答太泛。准备时应明确到具体键值、过期时间、读写路径、缓存失效策略，以及为什么不能只查数据库。",
)

add_question(
    doc,
    "多个骑手同时抢一个订单，如何避免重复接单？",
    "骑手抢单前会检查订单状态、支付状态、骑手字段和版本号等条件；只有数据库条件更新影响行数为 1 时才算抢单成功，否则说明订单已经被其他骑手拿走。",
    "这其实是项目里比较有价值的技术点，接近“条件更新＋乐观并发控制”的思路。但我现场解释得很绕，没有先讲竞态条件，再讲原子更新，最后讲成功判定。",
)

add_question(
    doc,
    "骑手端如何发现新订单？为什么这样轮询？",
    "骑手上线后会周期性刷新附近订单，我记得大约每 8 秒更新一次。至于轮询的具体实现和取舍，我当时没能讲清楚，只能承认这部分主要由 AI 完成。",
    "面试官继续追问的重点，很可能不是 8 秒这个数字，而是我是否理解定时轮询、请求重叠、页面卸载清理、前后台切换、失败重试，以及与 WebSocket 或服务器推送相比的优缺点。",
)

add_question(
    doc,
    "AI 写的代码，你怎么确认它真的能用？",
    "我主要从用户视角逐项手动测试功能。测试过程中确实发现并修复过问题，例如商家端发布新定价后用户端没有同步、腾讯地图页面闪跳，以及移动图钉后位置又回到原处。",
    "这些例子能证明我有质量意识，但还不足以证明工程可靠性。更完整的回答应包括代码审查、类型检查、单元测试、接口测试、端到端测试、日志和异常监控。",
)

doc.add_heading("3. 技术基础与 AI 编程：这是我被问穿的部分", level=2)
add_body(
    doc,
    "项目问完后，面试官把焦点放在我的基础能力上：是否真正写过前端代码、"
    "对 React 等框架了解多少、后端接触到什么程度，以及我有没有审查过 AI 生成的代码。"
    "我的回答比较诚实：前端系统学习很少，框架了解浅，代码实现大多依赖 AI，"
    "我更擅长拆需求、推动完成和手动验证。"
)
add_body(
    doc,
    "问题在于，诚实只能避免夸大，不能替代能力证明。"
    "当我说自己较信任模型、没有充分审查代码时，相当于直接暴露了项目的不可控风险。"
    "面试官真正关心的不是“能不能用 AI”，而是我能不能对 AI 的产出负责。"
)
add_callout(
    doc,
    "AI 可以替我加速实现，但不能替我承担解释、验证、调试和上线责任。",
    label="我听到的潜台词",
)

doc.add_heading("4. 个人优势与短板", level=2)
add_question(
    doc,
    "你觉得自己的强项是什么？",
    "我说自己抗压能力比较强、愿意拼，也有较强的执行力。别人给出任务后，我通常会马上去学、去做。",
    "这类回答最好配一个短例子，用行动证明“抗压”和“执行力”，否则很容易停留在自我评价。",
)
add_body(
    doc,
    "我也主动承认，自己在工程习惯上犯过低级错误，例如 Git 版本管理不规范、"
    "该提交时没有及时 commit。这个回答虽然真实，但如果只说错误、不说后来建立了什么机制，"
    "就会显得复盘没有闭环。"
)

doc.add_heading("5. 我的反向提问", level=2)
reverse_intro = add_body(doc, "最后我主要问了三个问题：")
reverse_intro.paragraph_format.keep_with_next = True
add_bullet(doc, "在 AI 时代，手写代码的能力是否仍然必要？", bullet_num_id)
add_bullet(doc, "前后端分别掌握到什么程度，才有机会进入团队实习？", bullet_num_id)
add_bullet(doc, "团队对算法题／LeetCode 的考察多不多？财经业务具体覆盖哪些方向？", bullet_num_id)
add_body(
    doc,
    "从对话中我的理解是：至少在当前求职和实习阶段，手写代码与独立理解代码仍然是基本要求。"
    "AI 能提升效率，但候选人必须具备足够的前后端基础，才能判断结果是否正确并持续推进复杂任务。"
    "财经业务则可以理解为覆盖支付、消费金融、保险等相关场景。"
)

doc.add_heading("三、这场面试暴露出的 5 个问题", level=1)
add_numbered(
    doc,
    "项目介绍没有主线。技术栈、业务流程和个人职责混在一起，导致面试官需要不断帮我“捞重点”。",
    decimal_num_id,
)
add_numbered(
    doc,
    "知道系统“做了什么”，但说不清“怎么做、为什么这样做”。部署、轮询和缓存都是典型例子。",
    decimal_num_id,
)
add_numbered(
    doc,
    "过度依赖 AI，缺少代码审查和自动化测试。项目能跑更多是结果证明，不是过程可控。",
    decimal_num_id,
)
add_numbered(
    doc,
    "技术概念缺少准确命名。并发抢单其实可以讲成条件更新、原子性和乐观并发控制，但我当时没有组织出来。",
    decimal_num_id,
)
add_numbered(
    doc,
    "工程基本功薄弱。Git、部署、日志、测试、框架基础都需要补齐，否则项目越大，风险越高。",
    decimal_num_id,
)

doc.add_heading("四、如果再来一次，我会这样准备", level=1)
add_numbered(
    doc,
    "把项目画成一张架构图，并能在 3 分钟内讲清：用户是谁、解决什么问题、系统怎么组成、我负责什么、最难的两个点是什么。",
    decimal_prepare_num_id,
)
add_numbered(
    doc,
    "对简历上的每个技术名词准备三层回答：它是什么、我为什么用、出了问题怎么排查。",
    decimal_prepare_num_id,
)
add_numbered(
    doc,
    "亲自重写或精读一个关键链路，例如“用户下单—商家接单—骑手抢单—状态流转”，确保能解释每次数据库更新与异常分支。",
    decimal_prepare_num_id,
)
add_numbered(
    doc,
    "补齐 AI 代码的质量闭环：查看 diff、做类型检查、写关键测试、记录日志、复现 bug、验证修复，而不是只看页面能否点通。",
    decimal_prepare_num_id,
)
add_numbered(
    doc,
    "重新整理反向提问，少问泛泛的“要学到什么程度”，多问团队真实的工作方式、实习生前 1—2 个月的任务、代码评审标准和评价机制。",
    decimal_prepare_num_id,
)

doc.add_heading("五、给非科班、AI 辅助开发同学的建议", level=1)
add_body(
    doc,
    "如果你和我一样，靠 AI 快速做出了一个完整产品，这件事本身很有价值。"
    "它证明你有需求理解、工具使用和推进落地的能力。但把项目写进简历之前，"
    "最好把“所有可能被追问的债”先还一遍：架构、数据流、并发、缓存、部署、测试、安全和故障处理。"
)
add_body(
    doc,
    "面试官通常不会因为你用了 AI 就否定项目，真正危险的是："
    "你把 AI 的产出当成了自己的能力，却无法解释关键代码，也无法证明系统可靠。"
    "最稳妥的做法，是把 AI 当作速度杠杆，把基础知识、判断力和责任感牢牢握在自己手里。"
)
add_callout(
    doc,
    "这次面试让我看清了下一阶段的方向：不是停止使用 AI，而是让自己有能力审查 AI、纠正 AI，并为最终结果负责。",
    label="写在最后",
)

doc.core_properties.title = "非科班大三，第一次被项目细节问穿｜字节跳动财经业务 AI 全栈实习面经"
doc.core_properties.subject = "面试经验帖"
doc.core_properties.keywords = "字节跳动, 财经业务, AI全栈, 实习, 面试复盘"
doc.core_properties.author = "匿名候选人"
doc.core_properties.last_modified_by = "匿名候选人"

doc.save(OUTPUT)
print(OUTPUT)
